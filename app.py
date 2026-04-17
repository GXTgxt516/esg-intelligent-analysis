import streamlit as st
import fitz
import os
import requests
import json
from io import BytesIO
from docx import Document

# --------------------------
# 页面配置
# --------------------------
st.set_page_config(
    page_title="ESG智能解析系统",
    page_icon="🌱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --------------------------
# 美化 CSS
# --------------------------
st.markdown("""
<style>
.main-title {
    font-size: 2.8rem !important;
    font-weight: 700;
    color: #0e4b70;
    text-align: center;
    margin-bottom: 1rem;
}
.box {
    background: #f7f9fc;
    padding: 20px;
    border-radius: 12px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    margin-bottom: 20px;
}
.question-btn {
    width: 100%;
    margin: 4px 0;
    border-radius: 8px;
}
</style>
""", unsafe_allow_html=True)

# --------------------------
# 会话状态初始化
# --------------------------
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "pdf_text" not in st.session_state:
    st.session_state.pdf_text = ""
if "auto_questions" not in st.session_state:
    st.session_state.auto_questions = []
if "trigger_q" not in st.session_state:
    st.session_state.trigger_q = None

API_KEY = os.getenv("API_KEY")

# --------------------------
# 侧边栏（系统菜单）
# --------------------------
with st.sidebar:
    st.title("🌱 系统菜单")
    st.divider()

    st.subheader("📌 功能导航")
    st.markdown("""
    - 上传 PDF 开始分析
    - 自动生成高频问题
    - 连续对话问答
    - 历史记录永久保存
    - 一键导出对话记录
    """)

    st.divider()
    st.subheader("⚙️ 对话管理")
    if st.button("🗑️ 清空所有对话", use_container_width=True):
        st.session_state.chat_history = []
        st.session_state.pdf_text = ""
        st.session_state.auto_questions = []
        st.rerun()

    if st.session_state.chat_history:
        if st.button("📥 导出对话为Word", use_container_width=True):
            doc = Document()
            doc.add_heading("ESG智能解析对话记录", 0)
            for item in st.session_state.chat_history:
                doc.add_heading("用户问题", level=2)
                doc.add_paragraph(item["user"])
                doc.add_heading("AI 回答", level=2)
                doc.add_paragraph(item["ai"])
                doc.add_page_break()
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            st.download_button(
                "点击下载Word",
                data=bio,
                file_name="ESG对话记录.docx",
                use_container_width=True
            )

    st.divider()
    st.caption("✅ ESG智能解析系统 | 修复版")

# --------------------------
# 主页面标题
# --------------------------
st.markdown('<div class="main-title">🌱 ESG智能报告解析系统</div>', unsafe_allow_html=True)

# --------------------------
# 第一步：上传文件（居中）
# --------------------------
st.markdown('<div class="box">', unsafe_allow_html=True)
st.subheader("📁 第一步：上传 ESG 报告 PDF")
uploaded_file = st.file_uploader("将文件拖入此处或点击选择", type="pdf")

if uploaded_file is not None:
    with st.spinner("正在解析PDF并生成高频问题..."):
        try:
            # 1. 解析PDF
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            # 限制文本长度，避免超出API限制
            st.session_state.pdf_text = text[:8000]

            # 2. 调用API生成高频问题
            prompt_q = f"""
            你是专业的ESG分析师，请根据以下ESG报告内容，生成8个用户最可能提问的高频问题，
            只返回问题列表，不要多余解释，也不要加序号。
            报告内容：
            {st.session_state.pdf_text}
            """

            headers = {
                "Authorization": f"Bearer {API_KEY}",
                "Content-Type": "application/json"
            }

            data = {
                "model": "qwen-turbo",
                "messages": [{"role": "user", "content": prompt_q}]
            }

            resp = requests.post(
                "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation",
                headers=headers,
                data=json.dumps(data),
                timeout=30
            )

            if resp.status_code == 200:
                res_json = resp.json()
                if "output" in res_json and "text" in res_json["output"]:
                    qs_text = res_json["output"]["text"].strip()
                    # 过滤空行和无效问题
                    qs_list = [q.strip() for q in qs_text.split("\n") if q.strip() and len(q.strip()) > 5]
                    if qs_list:
                        st.session_state.auto_questions = qs_list[:8]
                        st.success("✅ 解析完成！高频问题已自动生成")
                    else:
                        st.session_state.auto_questions = []
                        st.warning("⚠️ 解析成功，但未生成有效问题，请重试")
                else:
                    st.session_state.auto_questions = []
                    st.error("❌ API返回格式异常，未生成问题")
            else:
                st.session_state.auto_questions = []
                st.error(f"❌ API调用失败，状态码：{resp.status_code}")

        except Exception as e:
            st.session_state.pdf_text = ""
            st.session_state.auto_questions = []
            st.error(f"❌ 处理PDF时出错：{str(e)}")
st.markdown('</div>', unsafe_allow_html=True)

# --------------------------
# 第二步：显示高频问题（修复乱跑到对话区的问题）
# --------------------------
st.markdown('<div class="box">', unsafe_allow_html=True)
st.subheader("📋 第二步：系统自动生成高频问题（点击直接提问）")

if st.session_state.auto_questions:
    cols = st.columns(2)
    for i, q in enumerate(st.session_state.auto_questions):
        with cols[i % 2]:
            if st.button(q, key=f"q_btn_{i}", use_container_width=True):
                st.session_state.trigger_q = q
else:
    if uploaded_file is None:
        st.info("💡 上传PDF后，系统会自动生成高频问题")
    else:
        st.warning("⚠️ 未生成高频问题，请检查API Key或重新上传PDF")
st.markdown('</div>', unsafe_allow_html=True)

# --------------------------
# 第三步：连续智能问答
# --------------------------
st.markdown('<div class="box">', unsafe_allow_html=True)
st.subheader("💬 第三步：连续智能问答")

# 显示历史对话
for chat in st.session_state.chat_history:
    with st.chat_message("user"):
        st.write(chat["user"])
    with st.chat_message("assistant"):
        st.write(chat["ai"])

# 输入框
user_input = st.chat_input("输入你的问题……")

# 处理点击高频问题触发的提问
if st.session_state.trigger_q is not None:
    user_input = st.session_state.trigger_q
    st.session_state.trigger_q = None

if user_input:
    if not st.session_state.pdf_text:
        st.warning("⚠️ 请先上传PDF文件！")
    else:
        with st.chat_message("user"):
            st.write(user_input)

        prompt = f"""
        你是专业的ESG报告分析师，请根据以下报告内容和历史对话，回答用户的问题，回答要专业、简洁、有条理。
        报告内容：
        {st.session_state.pdf_text}

        历史对话：
        {st.session_state.chat_history}

        用户当前问题：{user_input}
        """

        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }

        data = {
            "model": "qwen-turbo",
            "messages": [{"role": "user", "content": prompt}]
        }

        with st.spinner("AI正在分析回答..."):
            try:
                response = requests.post(
                    "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation",
                    headers=headers,
                    data=json.dumps(data),
                    timeout=30
                )

                if response.status_code == 200:
                    res_json = response.json()
                    if "output" in res_json and "text" in res_json["output"]:
                        answer = res_json["output"]["text"]
                    else:
                        answer = "❌ API返回格式异常"
                else:
                    answer = f"❌ API调用失败，状态码：{response.status_code}"

            except Exception as e:
                answer = f"❌ 调用出错：{str(e)}"

        with st.chat_message("assistant"):
            st.write(answer)

        # 保存对话历史
        st.session_state.chat_history.append({
            "user": user_input,
            "ai": answer
        })

        # 刷新页面显示新对话
        st.rerun()
st.markdown('</div>', unsafe_allow_html=True)
